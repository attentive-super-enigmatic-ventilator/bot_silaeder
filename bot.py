import os, os.path
import ssl
import glob
import random
from random import random as rd
from hashlib import sha1
import urllib.request
from sys import stderr
from traceback import format_exc
import mimetypes
import smtplib
from email import encoders
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.audio import MIMEAudio
from email.mime.multipart import MIMEMultipart
import json
import pickle

import vk
import vk_api
from vk_api.longpoll import VkLongPoll
from vk_api.keyboard import VkKeyboard, VkKeyboardColor
from docx import Document
from docx.shared import Inches
import face_recognition
import requests

import downloading_google_sheet

user_service_access_key = open("user_service_access_key.dat").read().splitlines()[4]
service_access_key = open('group_service_access_key.dat').read().splitlines()
session = vk.Session(access_token=service_access_key)
api1 = vk.API(session, v=5.101)
owner_id = open('owner.dat').read().splitlines()
admins = [str(i) for i in open('admins.dat').read().splitlines()]
loginpassword = open('login_password.dat').read().splitlines()
mail1 = open('mail.dat').read().splitlines()
tokens = open('group_token.dat').read().splitlines()
vk_session = vk_api.VkApi(token=tokens[0])
vko = vk_session.get_api()
app_id = open('app.dat').read().splitlines()
sessio = vk.Session(access_token=user_service_access_key)
api = vk.API(sessio, v=5.101)
contacts = []
vk_sessio = vk_api.VkApi(loginpassword[0], loginpassword[1], app_id=int(app_id[0]), scope='wall, photos')
vk_sessio.auth()
upload = vk_api.VkUpload(vk_sessio)
admin_mails = [str(i) for i in open('admin_mail.dat').read().splitlines()]


known_faces_encodings = pickle.load(open("photos_prepared_for_face_recognition.dat", "rb"))
test_image_encodings = []
quantity_faces = []

users_report_filename = {}


def sendmail(text1, files):
    mail = smtplib.SMTP('smtp.mail.ru', 587)
    msg = MIMEMultipart()
    msg['From'] = mail1[0]
    msg['Subject'] = 'Новости Силаэдра'
    msg.attach(MIMEText(text1, 'plain'))
    for filepath in files:
        filename = os.path.basename(filepath)
        if os.path.isfile(filepath):
            ctype, encoding = mimetypes.guess_type(filepath)
            if ctype is None or encoding is not None:
                ctype = 'application/octet-stream'
            maintype, subtype = ctype.split('/', 1)
            if maintype == 'text':
                with open(filepath) as fp:
                    file = MIMEText(fp.read(), _subtype=subtype)
                    fp.close()
            elif maintype == 'image':
                with open(filepath, 'rb') as fp:
                    file = MIMEImage(fp.read(), _subtype=subtype)
                    fp.close()
            elif maintype == 'audio':
                with open(filepath, 'rb') as fp:
                    file = MIMEAudio(fp.read(), _subtype=subtype)
                    fp.close()
            else:
                with open(filepath, 'rb') as fp:
                    file = MIMEBase(maintype, subtype)
                    file.set_payload(fp.read())
                    fp.close()
                encoders.encode_base64(file)
            file.add_header('Content-Disposition', 'attachment', filename=filename)
            msg.attach(file)
    mail.starttls(context=ssl.create_default_context())
    mail.login(str(msg['From']), str(mail1[1]))
    mail.send_message(msg, to_addrs=contacts)
    mail.quit()


flag = False

base = VkKeyboard(one_time=True)
base.add_button('Отправить', color=VkKeyboardColor.POSITIVE)
base.add_button('Служебная записка')
base.add_button('Добавить нового человека', color=VkKeyboardColor.POSITIVE)
base.add_line()
base.add_button('help', color=VkKeyboardColor.POSITIVE)
base = base.get_keyboard()

back = VkKeyboard(one_time=True)
back.add_button('Вернуться к другим функциям', color=VkKeyboardColor.NEGATIVE)
back = back.get_keyboard()

def create_keyb1(buttons):
    keyboard = VkKeyboard(one_time=True)
    sh = 0
    for b in buttons:
        sh += 1
        if sh == 4:
            sh = 0
            keyboard.add_line()
        if b == 'new_line':
            keyboard.add_line()
        else:
            keyboard.add_button(b, color=VkKeyboardColor.POSITIVE)

    keyboard = keyboard.get_keyboard()
    return keyboard


def create_keyb(buttons):
    keyboard = VkKeyboard(one_time=True)
    sh = 0
    for b in buttons:
        sh += 1
        if sh == 4:
            sh = 0
            keyboard.add_line()
        if b == 'new_line':
            keyboard.add_line()
        else:
            keyboard.add_button(b, color=VkKeyboardColor.POSITIVE)
    keyboard.add_line()
    keyboard.add_button('Отменить отправку', color=VkKeyboardColor.NEGATIVE)
    keyboard = keyboard.get_keyboard()
    return keyboard


def create_keyb2(buttons):
    keyboard = VkKeyboard(one_time=True)
    for b in buttons:
        if b == 'new_line':
            keyboard.add_line()
        else:
            keyboard.add_button(b, color=VkKeyboardColor.POSITIVE)
    keyboard.add_line()
    keyboard.add_button('Отмена', color=VkKeyboardColor.NEGATIVE)
    keyboard = keyboard.get_keyboard()
    return keyboard

folder = 'photos'
for the_file in os.listdir(folder):
    file_path = os.path.join(folder, the_file)
    try:
        if os.path.isfile(file_path):
            os.unlink(file_path)
    except Exception as e:
        print(e, file=stderr)

stop = VkKeyboard(one_time=True)
stop.add_button('Отменить отправку', color=VkKeyboardColor.NEGATIVE)
stop = stop.get_keyboard()

f_mail = False
f_group = False
news = ''
users = {}
check_flag = False
check = ''
print("Successfully started", file=stderr)
while True:
    try:
        one_more_flag = False
        for event in vk_api.longpoll.VkLongPoll(vk_session).listen():
            if event.type == vk_api.longpoll.VkEventType.MESSAGE_NEW and event.to_me and str(event.user_id) in admins:
                if event.user_id not in users:
                    users[event.user_id] = 0
                incorrect_command = True
                inf = (vko.users.get(user_ids=event.user_id)[0])
                text = event.text.lower()

                if text == 'отменить отправку':
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Отправка отменена! Обращайтесь, когда появятся новости!',
                                      keyboard=base)
                    check = ''
                    check_flag = False
                    news = ''
                    incorrect_command = False
                    f_group = False
                    f_mail = False
                    users[event.user_id] = 0
                    contacts = []
                    for file in os.scandir():
                        if file.name.endswith(".docx"):
                            os.unlink(file.path)
                    folder = 'photos/' + str(event.user_id)
                    for the_file in os.listdir(folder):
                        file_path = os.path.join(folder, the_file)
                        try:
                            if os.path.isfile(file_path):
                                os.unlink(file_path)
                        except Exception as e:
                            print(e, file=stderr)
                    continue


                if text == 'нет, спасибо':
                    f_group = False
                    f_mail = False
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Хорошо! Обращайтесь, когда появятся еще новости!',
                                      keyboard=base)
                    users[event.user_id] = 0
                    if len(contacts) != 0:
                        sendmail(news, glob.glob("photos/" + str(event.user_id) + "/*.jpg"))
                    contacts = []
                    news = ''
                    incorrect_command = False
                    folder = 'photos/' + str(event.user_id)
                    for the_file in os.listdir(folder):
                        file_path = os.path.join(folder, the_file)
                        try:
                            if os.path.isfile(file_path):
                                os.unlink(file_path)
                        except Exception as e:
                            print(e, file=stderr)
                if text == 'отправить ранее выбранным контактам' or text == 'отправить новость':
                    check_flag = True
                    check = ''
                    for sy in range(5):
                        check += random.choice(['1', '2', '3', '4', '5', '6', '7', '8', '9', '0'])
                    mail = smtplib.SMTP('smtp.mail.ru', 587)
                    msg = MIMEMultipart()
                    msg['From'] = mail1[0]
                    msg['Subject'] = 'Код подтверждения'
                    msg.attach(MIMEText(check, 'plain'))
                    mail.starttls(context=ssl.create_default_context())
                    mail.login(str(msg['From']), str(mail1[1]))
                    mail.send_message(msg, to_addrs=admin_mails[admins.index(str(event.user_id))])
                    mail.quit()
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Я отправил Вам на почту ключ подтверждения. Отправьте его мне, чтобы успешно завершить рассылку',
                                      keyboard=stop)
                    incorrect_command = False
                    continue
                if check_flag and text == check:
                    check_flag = False
                    if not f_group:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Новость отправлена на почту выбранным контактам!' +
                                                  ' Вы хотите отправить новость еще куда-нибудь?',
                                          keyboard=create_keyb1(['Группа ВК', 'new_line', 'Нет, спасибо']))
                        f_mail = True
                        users[event.user_id] = 2
                        if len(contacts) != 0:
                            sendmail(news, glob.glob("photos/" + str(event.user_id) + "/*.jpg"))
                        check_flag = False
                        continue
                    else:
                        f_group = False
                        f_mail = False
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Новость отправлена на почту выбранным контактам! ' +
                                                  'Обращайтесь, когда появятся еще новости!',
                                          keyboard=base)
                        users[event.user_id] = 0
                        if len(contacts) != 0:
                            sendmail(news, glob.glob("photos/" + str(event.user_id) + "/*.jpg"))
                        contacts = []
                        news = ''
                        incorrect_command = False
                        folder = 'photos/' + str(event.user_id)
                        for the_file in os.listdir(folder):
                            file_path = os.path.join(folder, the_file)
                            try:
                                if os.path.isfile(file_path):
                                    os.unlink(file_path)
                            except Exception as e:
                                print(e, file=stderr)
                    incorrect_command = False
                    check_flag = False
                    continue
                if users[event.user_id] == 3 and check_flag and text != check:
                    check_flag = True
                    check = ''
                    for sy in range(5):
                        check += random.choice(['1', '2', '3', '4', '5', '6', '7', '8', '9', '0'])
                    mail = smtplib.SMTP('smtp.mail.ru', 587)
                    msg = MIMEMultipart()
                    msg['From'] = mail1[0]
                    msg['Subject'] = 'Код подтверждения'
                    msg.attach(MIMEText(check, 'plain'))
                    mail.starttls(context=ssl.create_default_context())
                    mail.login(str(msg['From']), str(mail1[1]))
                    mail.send_message(msg, to_addrs=admin_mails[admins.index(str(event.user_id))])
                    mail.quit()
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Неправильный код подтверждения!!! Я отправил Вам на почту новый ключ подтверждения. Отправьте его мне, чтобы успешно завершить рассылку',
                                      keyboard=stop)
                    incorrect_command = False
                    continue
                if users[event.user_id] == 1 and text == 'нет':
                    users[event.user_id] = 2
                    one_more_flag = True
                    incorrect_command = False
                if users[event.user_id] == 1 and text == 'да':
                    downloading_google_sheet.auth()
                    downloading_google_sheet.send_to_some(firstname + ' ' + lastname)
                    contacts = downloading_google_sheet.contacts
                    downloading_google_sheet.clear()
                    try:
                        sendmail(news, glob.glob("photos/" + str(event.user_id) + "/*.jpg"))
                    except Exception as e:
                        print('ОШИБКА: ' + str(e), file=stderr)
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Новость отправлена! Обращайтесь, когда появятся еще новости!',
                                      keyboard=base)
                    users[event.user_id] = 0
                    contacts = []
                    news = ''
                    incorrect_command = False
                    folder = 'photos/' + str(event.user_id)
                    for the_file in os.listdir(folder):
                        file_path = os.path.join(folder, the_file)
                        try:
                            if os.path.isfile(file_path):
                                os.unlink(file_path)
                        except Exception as e:
                            print(e, file=stderr)
                    continue
                if users[event.user_id] == 1:
                    result = {}
                    news = event.text
                    flag = False
                    photos = api1.messages.getById(message_ids=event.message_id, group_id=183112747)
                    try:
                        os.mkdir("photos/" + str(event.user_id))
                    except:
                        pass
                    for i in range(len(photos['items'][0]['attachments'])):

                        if photos['items'][0]['attachments'][i]['type'] == 'photo':
                            length = len(photos['items'][0]['attachments'][i]['photo']['sizes']) - 1
                            urllib.request.urlretrieve(
                                photos['items'][0]['attachments'][i]['photo']['sizes'][length]['url'],
                                'photos/' + str(event.user_id) + "/" + str(i) + '.jpg')

                        if photos['items'][0]['attachments'][i]['type']== 'doc':
                            urllib.request.urlretrieve(
                                photos['items'][0]['attachments'][i]['doc']['url'],
                                'photos/' + str(event.user_id) + "/" + str(i) + '.' + photos['items'][0]['attachments'][i]['doc']['ext'])
                    known_faces_encodings = pickle.load(
                        open("photos_prepared_for_face_recognition.dat", "rb"))
                    face = False
                    test_image_encodings = []
                    quantity_faces = []
                    for i in glob.glob("photos/" + str(event.user_id) + "/*.jpg"):
                        test_image = face_recognition.load_image_file(i)
                        test_image_locations = face_recognition.face_locations(test_image)
                        if not test_image_locations == []:
                            test_image_encoding = face_recognition.face_encodings(test_image, test_image_locations)[
                                0]
                            test_image_encodings.append(test_image_encoding)
                            quantity_faces.append(test_image_locations)
                    f = False
                    if not quantity_faces == []:
                        for j in test_image_encodings:
                            result = {}
                            for i in known_faces_encodings.keys():
                                comparing = face_recognition.compare_faces([known_faces_encodings[i]], j)[0]

                                if comparing:
                                    name = i.split('/')[-1].split('_')
                                    name[1] = name[1][0:(len(name[1]) - 4)]
                                    try:
                                        d = int(name[1][-1])
                                        name[1] = name[1][:len(name[1] - 1)]
                                    except:
                                        pass
                                    try:
                                        result[name[0] + ' ' + name[1]] += 1
                                    except:
                                        result[name[0] + ' ' + name[1]] = 1
                                    face = True
                            if face:
                                maximum = 0
                                for key in result:
                                    if result[key] > maximum:
                                        maximum = result[key]
                                        firstname, lastname = key.split()
                                f = True
                                vko.messages.send(user_id=event.user_id,
                                                  random_id=random.randint(1, 10 ** 9),
                                                  message='Вы хотите отправить это на почту родителям ребенка, который был найден на этой фотографии(' + firstname + ' ' + lastname + ')', keyboard=create_keyb1(['Да', 'Нет']))

                    if not f:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Где Вы хотите опубликовать новость?',
                                          keyboard=create_keyb(['Группа ВК', 'Почта']))

                        users[event.user_id] = 2
                        continue
                    incorrect_command = False
                if text == 'отправить':
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Напишите мне, что Вы хотите разослать. ' +
                                              'При необходимости Вы можете прикрепить файлы к сообщению.',
                                      keyboard=stop)
                    flag = True
                    users[event.user_id] = 1
                    incorrect_command = False
                if 'привет' in text and users[event.user_id] == 0:
                    if vko.users.get(user_ids=event.user_id, fields=['sex'])[0]['sex'] == 2:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Привет, ' + (inf['first_name']) + ' ' +
                                                  inf['last_name'] + '!' + ' Я рад, что ты мне написал!',
                                          keyboard=base)
                    else:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Привет, ' + (inf['first_name']) + ' ' +
                                                  inf['last_name'] + '!' + ' Я рад, что ты мне написала!',
                                          keyboard=base)
                    incorrect_command = False
                    continue
                if users[event.user_id] == 10:
                    x = True
                    if 'родителям' in text:
                        downloading_google_sheet.send_to_class(grade)
                        contacts += downloading_google_sheet.contacts
                        downloading_google_sheet.clear()
                        users[event.user_id] = 3
                    elif 'ученикам' in text:
                        downloading_google_sheet.send_to_children(grade)
                        contacts += downloading_google_sheet.contacts
                        downloading_google_sheet.clear()
                        users[event.user_id] = 3
                    elif 'всем' in text:
                        downloading_google_sheet.send_to_class(grade)
                        downloading_google_sheet.send_to_children(grade)
                        contacts += downloading_google_sheet.contacts
                        downloading_google_sheet.clear()
                        users[event.user_id] = 3
                    else:
                        x = False
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Я Вас не понимаю, используйте кнопки ниже!',
                                          keyboard=create_keyb(['Родителям(' + grade + ')', 'Ученикам(' + grade + ')',
                                                                'Всем(' + grade + ')']))
                        incorrect_command = False
                    if x:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Вы хотите отправить новость еще кому-нибудь? Если хотите, то напишите кому еще' +
                                                  ' надо отправить новость, а если хотите отправить новость ранее выбранным' +
                                                  ' контактам, то нажмите на кнопку "Отправить ранее выбранным контактам".' +
                                                  '\n' + 'Для отмены отправки нажмите на кнопку "Отменить отправку"',
                                          keyboard=create_keyb(
                                              downloading_google_sheet.auth() + ['new_line', 'Учителям',
                                               'Отправить ранее выбранным контактам']))
                        continue
                if users[event.user_id] == 3:
                    cont = False
                    to_all = False
                    incorrect_command = False
                    grade = '0'
                    downloading_google_sheet.auth()
                    for i in range(len(text)):
                        try:
                            grade = str(int(text[i]))
                            try:
                                a = str(int(text[i + 1]))
                                grade += a
                                grade += ' ' + text[i + 3].upper()
                            except:
                                grade += text[i + 2].upper()
                            break
                        except:
                            pass
                    number_of_contacts = len(contacts)
                    if grade != '0':
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Кому Вы хотите отправить новость?',
                                          keyboard=create_keyb(['Родителям(' + grade + ')', 'Ученикам(' + grade + ')',
                                                                'Всем(' + grade + ')']))
                        cont = True
                        users[event.user_id] = 10
                        continue
                    elif 'абсолютно всем' in text:
                        downloading_google_sheet.send_to_all()
                        contacts += downloading_google_sheet.contacts
                        downloading_google_sheet.clear()
                        to_all = True
                    elif 'учителям' in text:
                        downloading_google_sheet.send_to_teachers()
                        contacts += downloading_google_sheet.contacts
                        downloading_google_sheet.clear()
                    else:
                        downloading_google_sheet.send_to_some(event.text)
                        contacts += downloading_google_sheet.contacts
                        downloading_google_sheet.clear()
                    if len(contacts) == number_of_contacts:
                        cont = True
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Я Вас не понимаю, используйте кнопки ниже!',
                                          keyboard=create_keyb(
                                              downloading_google_sheet.auth() + ['new_line',
                                               'Учителям', 'Абсолютно всем']))
                        continue
                    if to_all:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Подтвердите отправку', keyboard=create_keyb(['Отправить новость']))
                    elif not cont:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Вы хотите отправить новость еще кому-нибудь? Если хотите, то ' +
                                                  'напишите кому еще надо отправить новость, а если хотите отправить новость ранее ' +
                                                  'выбранным контактам, то нажмите на кнопку "Отправить ранее выбранным контактам".' +
                                                  '\n' + 'Для отмены отправки нажмите на кнопку "Отменить отправку"',
                                          keyboard=create_keyb(
                                              downloading_google_sheet.auth() + ['new_line',
                                               'Учителям',
                                               'Отправить ранее выбранным контактам']))
                if text == 'почта' and users[event.user_id] == 2 or one_more_flag:
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Кому Вы хотите отправить новость? Вы можете отправить новость всем родителям ' +
                                              'определенного класса с помощью кнопок снизу.' + '\n' +
                                              'Также Вы можете отправить новость конкретным родителям, написав сообщение типа' +
                                              '"маме <Имя ребенка> <Фамилия ребенка>" или ' +
                                              '"отцу <Имя ребенка> <Фамилия ребенка>".',
                                      keyboard=create_keyb(
                                          downloading_google_sheet.auth() + ['new_line',
                                           'Учителям', 'Абсолютно всем']))
                    users[event.user_id] = 3
                    incorrect_command = False
                    one_more_flag = False
                    continue
                if text == 'опубликовать новость' and users[event.user_id] == 1001:
                    photos = glob.glob("photos/" + str(event.user_id) + "/*.jpg")
                    if len(photos) != 0:
                        photo_list = upload.photo_wall(photos)
                        attachment = ','.join('photo{owner_id}_{id}'.format(**item) for item in photo_list)
                        api.wall.post(owner_id=owner_id, message=news, attachments=attachment)
                    else:
                        api.wall.post(owner_id=owner_id, message=news)
                    f_group = True
                    if not f_mail:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Новость выложена в группе Силаэдра в ВК! Вы хотите отправить ее еще ' +
                                                  'куда-нибудь?',
                                          keyboard=create_keyb1(['Почта', 'new_line', 'Нет, спасибо']))
                    else:
                        f_group = False
                        f_mail = False
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Новость выложена в группе Силаэдра в ВК! Обращайтесь, когда появятся еще' +
                                                  ' новости!',
                                          keyboard=base)
                        users[event.user_id] = 0
                        if len(contacts) != 0:
                            sendmail(news, glob.glob("photos/" + str(event.user_id) + "/*.jpg"))
                        contacts = []
                        news = ''
                        incorrect_command = False
                        folder = 'photos/' + str(event.user_id)
                        for the_file in os.listdir(folder):
                            file_path = os.path.join(folder, the_file)
                            try:
                                if os.path.isfile(file_path):
                                    os.unlink(file_path)
                            except Exception as e:
                                print(e, file=stderr)
                    incorrect_command = False
                    continue
                if text == 'группа вк' and users[event.user_id] == 2:
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Подтвердите публикацию новости',
                                      keyboard=create_keyb(['Опубликовать новость']))
                    users[event.user_id] = 1001
                    incorrect_command = False
                    continue
                if text == 'сайт' and users[event.user_id] == 2:
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='К сожалению, сайт Силаэдра пока не работает! Вы хотите отправить новость ' +
                                              'еще куда-нибудь?',
                                      keyboard=create_keyb1(['Группа ВК', 'Сайт', 'Почта', 'new_line', 'Нет, спасибо']))

                    incorrect_command = False
                    continue
                if users[event.user_id] == 2:
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Я Вас не понимаю, используйте кнопки ниже.',
                                      keyboard=create_keyb(['Группа ВК', 'Почта']))
                    incorrect_command = False
                    continue
                if users[event.user_id] == 1001:
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Я Вас не понимаю, используйте кнопки ниже.',
                                      keyboard=create_keyb(['Опубликовать новость']))
                    incorrect_command = False
                    continue
                if text == 'help' or text == 'помощь' or text == '/start' or text == '/help' or text == 'начать':
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Мои функции:' + '\n' +
                                              '- напишите мне "Отправить" для автоматической рассылки новостей, далее ' +
                                              'следуйте моим указаниям' + '\n' +
                                              '- напишите мне "Служебная записка", чтобы приступить к составлению служебной записки',
                                      keyboard=base)
                    incorrect_command = False
                if text == 'вернуться к другим функциям':
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Вы вернулись в главное меню',
                                      keyboard=base)
                    news = ''
                    incorrect_command = False
                    f_group = False
                    f_mail = False
                    users[event.user_id] = 0
                    contacts = []
                    for file in os.scandir():
                        if file.name.endswith(".docx"):
                            os.unlink(file.path)
                    folder = 'photos/' + str(event.user_id)
                    for the_file in os.listdir(folder):
                        file_path = os.path.join(folder, the_file)
                        try:
                            if os.path.isfile(file_path):
                                os.unlink(file_path)
                        except Exception as e:
                            print(e, file=stderr)
                    continue
                if text == 'служебная записка':
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Здравствуйте, дорогой учитель! Это функция "Служебная записка". Вы должны ответить на несколько вопросов про это мероприятие, и бот пришлёт вам на почту документ.' + '\n' +
                                              '\n' + 'Что это за мероприятие? Отправьте мне текстовое сообщение', keyboard=back)
                    users[event.user_id] = 73
                    continue
                if users[event.user_id] == 73:
                    users_report_filename[event.user_id] = Document()
                    report_file_name = sha1(str(rd()).encode('utf-8')).hexdigest() + '.docx'
                    users_report_filename[event.user_id].add_paragraph('Директору').paragraph_format.left_indent = Inches(4.85)
                    users_report_filename[event.user_id].add_paragraph('ГБОУ “Школа им. Маршала В.И.Чуйкова').paragraph_format.left_indent = Inches(
                        2.89)
                    users_report_filename[event.user_id].add_paragraph("А.А.Инглези от О.А.Старуновой").paragraph_format.left_indent = Inches(3.47)
                    event_description = event.text
                    users_report_filename[event.user_id].add_paragraph(event_description).paragraph_format.left_indent = Inches(0.5)
                    users[event.user_id] = 74
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Когда будет проходить это мероприятие?', keyboard=back)
                    continue

                if users[event.user_id] == 74:
                    event_date = event.text
                    users_report_filename[event.user_id].add_paragraph('1)Дата мероприятия: ' + event_date)
                    users[event.user_id] = 75
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Где будет проходить это мероприятие?', keyboard=back)
                    continue

                if users[event.user_id] == 75:
                    event_place = event.text
                    users_report_filename[event.user_id].add_paragraph('2)Место проведения: ' + event_place)
                    users[event.user_id] = 76
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Кто будет участвовать в этом мероприятии?',
                                      keyboard=create_keyb1(
                                          downloading_google_sheet.auth()))

                    users_report_filename[event.user_id].add_paragraph('3)Участники: ' + '\n')
                    n = 1
                    table = users_report_filename[event.user_id].add_table(rows=n, cols=5)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Класс'
                    hdr_cells[1].text = 'Имя'
                    hdr_cells[2].text = 'Фамилия'
                    hdr_cells[3].text = 'Отчество'
                    hdr_cells[4].text = 'Дата рождения'
                    members = []
                    incorrect_command = False
                    continue

                if users[event.user_id] == 76 and text == 'все участники уже перечислены':
                    users[event.user_id] = 77
                    downloading_google_sheet.auth_teachers()
                    s = ''
                    for i in range(2, len(downloading_google_sheet.sheet1)):
                        s += str(i - 1) + ') ' + downloading_google_sheet.sheet1[i][0] + ' ' + \
                             downloading_google_sheet.sheet1[i][1] + ' ' + downloading_google_sheet.sheet1[i][2] + '\n'
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Кто из учителей будет участвовать в этом мероприятии? Перечислите через пробел номера учителей, которые будут участвоать в мероприятии.' + '\n' + s, keyboard=back)
                    incorrect_command = False
                    continue
                elif users[event.user_id] == 76:
                    downloading_google_sheet.auth()
                    fl = False
                    for i in range(len(downloading_google_sheet.sheet)):
                        if downloading_google_sheet.sheet[i][0] == event.text:
                            fl = True
                            row_cells = table.add_row().cells
                            row_cells[0].text = event.text
                            row_cells[1].text = downloading_google_sheet.sheet[i][2]
                            row_cells[2].text = downloading_google_sheet.sheet[i][1]
                            row_cells[3].text = downloading_google_sheet.sheet[i][3]
                            row_cells[4].text = downloading_google_sheet.sheet[i][4]
                        elif downloading_google_sheet.sheet[i][0] == '':
                            break
                    if fl:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Кто еще будет участвовать в этом мероприятии?',
                                          keyboard=create_keyb1(
                                              downloading_google_sheet.auth() +
                                                ['Все участники уже перечислены']))
                    else:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Используйте кнопки ниже, я Вас не понимаю!', keyboard=create_keyb1(
                                            downloading_google_sheet.auth() +
                                            ['Все участники уже перечислены']))
                    incorrect_command = False

                if users[event.user_id] == 77:
                    try:
                        numb = [int(i) for i in event.text.split()]
                        users_report_filename[event.user_id].add_paragraph('4)Учителя:')
                        k = 1
                        table = users_report_filename[event.user_id].add_table(rows=k, cols=5)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Номер'
                        hdr_cells[1].text = 'Имя'
                        hdr_cells[2].text = 'Фамилия'
                        hdr_cells[3].text = 'Отчество'
                        hdr_cells[4].text = 'Предмет'
                        for i in range(len(numb)):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(i + 1)
                            row_cells[1].text = downloading_google_sheet.sheet1[numb[i] + 1][1]
                            row_cells[2].text = downloading_google_sheet.sheet1[numb[i] + 1][0]
                            row_cells[3].text = downloading_google_sheet.sheet1[numb[i] + 1][2]
                            row_cells[4].text = downloading_google_sheet.sheet1[numb[i] + 1][3]
                        users[event.user_id] = 78
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Кто будет ответственным за это мероприятие?', keyboard=back)
                    except:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 10 ** 9),
                                          message='Неправильный формат ввода! Отправьте мне номера учителей через пробел.', keyboard=back)
                    incorrect_command = False
                    continue
                if users[event.user_id] == 78:
                    event_responsible = event.text
                    users_report_filename[event.user_id].add_paragraph('\n5)Ответственный: ' + event_responsible)
                    users_report_filename[event.user_id].save(report_file_name)
                    upload_url = vko.docs.getMessagesUploadServer(type='doc', peer_id='489061359')['upload_url']
                    response = requests.post(upload_url, files={'file': open(report_file_name, 'rb')})
                    result = json.loads(response.text)
                    file = result['file']
                    json1 = vko.docs.save(file=file, title=report_file_name, tags=[])

                    owner_id1 = json1['doc']['owner_id']
                    doc_id = json1['doc']['id']
                    attach = 'doc' + str(owner_id1) + '_' + str(doc_id)

                    vk_session = vk_api.VkApi(token=tokens[0])
                    vko = vk_session.get_api()
                    for sy in range(5):
                        check += random.choice(['1', '2', '3', '4', '5', '6', '7', '8', '9', '0'])
                    mail = smtplib.SMTP('smtp.mail.ru', 587)
                    msg = MIMEMultipart()
                    msg['From'] = mail1[0]
                    msg['Subject'] = 'Код подтверждения'
                    msg.attach(MIMEText(check, 'plain'))
                    mail.starttls(context=ssl.create_default_context())
                    mail.login(str(msg['From']), str(mail1[1]))
                    mail.send_message(msg, to_addrs=admin_mails[admins.index(str(event.user_id))])
                    mail.quit()
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Я отправил Вам на почту ключ подтверждения. Отправьте его мне, чтобы успешно завершить рассылку',
                                      keyboard=stop)
                    users[event.user_id] = 11111
                    continue
                if text == check and users[event.user_id] == 11111:

                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Держите!',
                                      attachment=attach, keyboard=base)

                    os.remove(report_file_name)
                    incorrect_command = False
                    users[event.user_id] = 0
                    continue
                elif text != check:
                    vk_session = vk_api.VkApi(token=tokens[0])
                    vko = vk_session.get_api()
                    for sy in range(5):
                        check += random.choice(['1', '2', '3', '4', '5', '6', '7', '8', '9', '0'])
                    mail = smtplib.SMTP('smtp.mail.ru', 587)
                    msg = MIMEMultipart()
                    msg['From'] = mail1[0]
                    msg['Subject'] = 'Код подтверждения'
                    msg.attach(MIMEText(check, 'plain'))
                    mail.starttls(context=ssl.create_default_context())
                    mail.login(str(msg['From']), str(mail1[1]))
                    mail.send_message(msg, to_addrs=admin_mails[admins.index(str(event.user_id))])
                    mail.quit()
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Неправильный код подтверждения, я отправил Вам на почту новый. Отправьте его мне, чтобы успешно завершить рассылку',
                                      keyboard=stop)
                    users[event.user_id] = 11111
                    continue
                if incorrect_command:
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 10 ** 9),
                                      message='Извините, я Вас не понимаю. ' + '\n' +
                                              'Мои функции:' + '\n' +
                                              '- напишите мне "Отправить" для автоматической рассылки новостей, ' +
                                              'далее следуйте моим указаниям' + '\n' +
                                              '- напишите мне "Привет", и я поздороваюсь с Вами', keyboard=base)
                    users[event.user_id] = 0
            elif event.type == vk_api.longpoll.VkEventType.MESSAGE_NEW and event.to_me and str(event.user_id) not in admins:
                if str(event.user_id) not in users:
                    users[str(event.user_id)] = 0
                text = event.text.lower()
                inf = (vko.users.get(user_ids=event.user_id)[0])
                if 'привет' in text or 'начать' in text:
                    if vko.users.get(user_ids=event.user_id, fields=['sex'])[0]['sex'] == 2:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 109),
                                          message = 'Привет, ' + (inf['first_name']) + ' ' + inf['last_name'] + '!' + ' Я рад, что ты мне написал!',
                                          keyboard=create_keyb1(['Узнать расписание']))
                    else:
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 109),
                                          message = 'Привет, ' + (inf['first_name']) + ' ' + inf['last_name'] + '!' + ' Я рад, что ты мне написала!',
                                          keyboard=create_keyb1(['Узнать расписание']))
                        incorrect_command = False
                if text == 'отмена':
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 109),
                                      message='Теперь вы снова в начале пути',
                                      keyboard=create_keyb1(['Узнать расписание']))

                if text == "узнать расписание":
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 109),
                                      message='Что вас интересует?',
                                      keyboard=create_keyb2(['Вся неделя', "Сегодня", "Завтра"]))

                if users[str(event.user_id)] == 1:
                    if text in downloading_google_sheet.auth():
                        sum = downloading_google_sheet.check_timetable(text=text)
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 109),
                                          message=sum,
                                          keyboard=create_keyb1(['Узнать расписание']))

                if users[str(event.user_id)] == 2:
                    if event.text in downloading_google_sheet.auth():
                        sum = downloading_google_sheet.today(text=text)
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 109),
                                          message=sum,
                                          keyboard=create_keyb1(['Узнать расписание']))

                if users[str(event.user_id)] == 3:
                    if event.text in downloading_google_sheet.auth():
                        sum = downloading_google_sheet.tomorrow(text=text)
                        vko.messages.send(user_id=event.user_id,
                                          random_id=random.randint(1, 109),
                                          message=sum,
                                          keyboard=create_keyb1(['Узнать расписание']))

                if text == "сегодня":
                    users[str(event.user_id)] = 2
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 109),
                                      message='Расписание для какого класса вас интересует?',
                                      keyboard=create_keyb2(downloading_google_sheet.classes_names))

                if text == "завтра":
                    users[str(event.user_id)] = 3
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 109),
                                      message='Расписание для какого класса вас интересует?',
                                      keyboard=create_keyb2(downloading_google_sheet.auth()))

                if text == "вся неделя":
                    users[str(event.user_id)] = 1
                    vko.messages.send(user_id=event.user_id,
                                      random_id=random.randint(1, 109),
                                      message='Расписание для какого класса вас интересует?',
                                      keyboard=create_keyb2(downloading_google_sheet.auth()))

    except Exception as er:
        print(er, file=stderr)
        users[event.user_id] = 0
        pass
